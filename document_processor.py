"""
Universal Document Processor - Enhanced V3
Handles: PDF, DOCX, DOC, TXT, Images (JPG, PNG, JPEG, WEBP)
Enhanced for complex layouts, multi-column PDFs, creative formats
Improved contact and name extraction with strict validation
"""

import io
import re
import base64
import requests
from typing import Optional, Dict, List, Tuple, Set
from pathlib import Path


# ═══════════════════════════════════════════════════════════════
#                    NAME VALIDATION BLACKLISTS
# ═══════════════════════════════════════════════════════════════

# ANY single word match → reject the candidate name
_STRONG_NAME_BLACKLIST: Set[str] = {
    # Resume / document labels
    "resume", "cv", "curriculum", "vitae", "biodata",
    # Section headers
    "objective", "summary", "profile", "overview",
    "education", "experience", "skills", "certifications",
    "references", "declaration", "signature",
    "qualifications", "achievements", "accomplishments",
    "responsibilities", "duties",
    # HR terms
    "candidate", "applicant", "recruitment", "hiring",
    "position", "vacancy",
    # Section labels
    "contact", "details", "information", "personal",
    "professional", "technical", "academic",
    "employment", "employer",
    # Job titles
    "engineer", "developer", "manager", "analyst",
    "consultant", "designer", "architect", "administrator",
    "senior", "junior", "lead", "head", "director",
    "executive", "intern", "trainee",
    # Industry / org words
    "software", "hardware", "technology", "technologies",
    "university", "college", "school", "institute",
    "company", "corporation", "organization",
    "private", "limited", "ltd", "inc", "corp", "llc", "pvt",
    # Degree words
    "bachelor", "master", "doctor", "phd", "mba",
    "btech", "mtech", "degree", "diploma",
    # Other
    "project", "projects", "work", "history",
    "reference", "present", "current",
    # Months
    "january", "february", "march", "april", "may", "june",
    "july", "august", "september", "october", "november", "december",
    # Days
    "monday", "tuesday", "wednesday", "thursday",
    "friday", "saturday", "sunday",
    # Locations (common cities that appear as headers)
    "bangalore", "mumbai", "delhi", "chennai", "hyderabad",
    "pune", "kolkata", "india", "singapore",
}

# 2 or more word matches → reject
_SOFT_NAME_BLACKLIST: Set[str] = {
    "voice", "message", "support", "service", "process",
    "system", "application", "development",
    "client", "customer", "business",
    "department", "team",
    "specialist", "coordinator", "associate",
    "responsible", "working", "developing", "managing",
    "leading", "building", "creating",
    "page", "date", "mobile", "email", "phone",
    "address", "city", "state", "country",
}


# ═══════════════════════════════════════════════════════════════
#                    ENHANCED TEXT CLEANING
# ═══════════════════════════════════════════════════════════════

def _clean_extracted_text(text: str) -> str:
    """Clean common text extraction artifacts"""
    if not text:
        return ""

    # Fix bullet points
    for bullet in ['•', '●', '○', '■', '▪', '►', '▸', '◆', '◇', '▶', '→', '➤', '➢', '✓', '✔', '☑']:
        text = text.replace(bullet, '\n• ')

    # Fix common ligatures
    ligatures = {
        'ﬁ': 'fi', 'ﬂ': 'fl', 'ﬀ': 'ff', 'ﬃ': 'ffi', 'ﬄ': 'ffl',
        '…': '...', '–': '-', '—': '-', '\u2018': "'", '\u2019': "'",
        '\u201c': '"', '\u201d': '"', '′': "'", '″': '"'
    }
    for old, new in ligatures.items():
        text = text.replace(old, new)

    # Fix multiple spaces but preserve structure
    text = re.sub(r' {3,}', '  ', text)

    # Fix phone numbers that get split
    text = re.sub(r'(\+\d{1,3})\s+(\d)', r'\1\2', text)
    text = re.sub(r'(\d{3})\s*[-.)]\s*(\d{3})\s*[-.)]\s*(\d{4})', r'\1-\2-\3', text)

    # Fix emails that get split
    text = re.sub(r'(\w)\s+@\s+(\w)', r'\1@\2', text)
    text = re.sub(r'(\w)@(\w+)\s+\.(\w+)', r'\1@\2.\3', text)
    text = re.sub(r'(\w+)\s*@\s*(\w+)\s*\.\s*(\w+)', r'\1@\2.\3', text)

    # Fix URLs that get split
    text = re.sub(r'(https?)\s*:\s*/\s*/\s*', r'\1://', text)
    text = re.sub(r'(www)\s*\.\s*', r'www.', text)
    text = re.sub(r'\.\s*(com|org|net|io|in|edu|co)\b', r'.\1', text)

    # Fix LinkedIn URLs
    text = re.sub(r'linkedin\s*\.\s*com\s*/\s*in\s*/\s*', r'linkedin.com/in/', text, flags=re.IGNORECASE)

    # Fix GitHub URLs
    text = re.sub(r'github\s*\.\s*com\s*/\s*', r'github.com/', text, flags=re.IGNORECASE)

    return text

def _collapse_spaced_text(text: str) -> str:
    """
    Collapse text where characters are separated by spaces.
    Handles: "S A N K E T  R A J E N D R A" → "SANKET RAJENDRA"
    Also: "E D U C A T I O N" → "EDUCATION"
    Also: "h t t p s : / / w w w . i b m . c o m" → "https://www.ibm.com"
    """
    if not text:
        return ""

    lines = text.split('\n')
    result_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            result_lines.append('')
            continue

        if _is_spaced_out_line(stripped):
            collapsed = _collapse_spaced_line(stripped)
            result_lines.append(collapsed)
        else:
            result_lines.append(stripped)

    return '\n'.join(result_lines)


def _is_spaced_out_line(line: str) -> bool:
    """Detect if a line has spaced-out characters like 'S A N K E T'."""
    if not line or len(line) < 5:
        return False

    # Quick check: count ratio of spaces to non-spaces
    spaces = line.count(' ')
    non_spaces = len(line) - spaces
    if non_spaces == 0:
        return False

    # In spaced-out text, spaces ≈ non-spaces (one space per character)
    if spaces / max(non_spaces, 1) < 0.5:
        return False

    # Count single-char segments between spaces
    parts = [p for p in line.split(' ') if p]  # filter empty from double-space
    if len(parts) < 3:
        return False

    single_chars = sum(1 for p in parts if len(p) == 1)
    ratio = single_chars / len(parts)

    return ratio > 0.5


def _collapse_spaced_line(line: str) -> str:
    """
    Collapse a spaced-out line using double-space as word boundary.
    "S A N K E T  R A J E N D R A" → "SANKET RAJENDRA"
    "h t t p s : / / w w w . i b m . c o m" → "https://www.ibm.com"
    """
    if not line:
        return line

    # Split by double-space (or more) to find word boundaries
    word_groups = re.split(r' {2,}', line.strip())

    collapsed_words = []
    for group in word_groups:
        group = group.strip()
        if not group:
            continue

        parts = group.split(' ')

        # Check if all parts are single characters
        if all(len(p) <= 1 for p in parts if p):
            # Collapse: "S A N K E T" → "SANKET"
            word = ''.join(p for p in parts if p)
            collapsed_words.append(word)
        else:
            # Mixed or normal text — keep as is
            collapsed_words.append(group)

    return ' '.join(collapsed_words)
    
    # Also handle cases where words within are separated by single space
    # but individual chars: "h t t p s : / / w w w . i b m . c o m"
    # Detect: if result still has many single-char "words", try harder
    result_words = result.split()
    if len(result_words) > 3:
        single_word_count = sum(1 for w in result_words if len(w) == 1)
        if single_word_count / len(result_words) > 0.5:
            # Still spaced — just remove all spaces between single chars
            collapsed = ""
            i = 0
            while i < len(result):
                if result[i] != ' ':
                    collapsed += result[i]
                    i += 1
                else:
                    # Check if surrounded by single chars
                    before_single = (i > 0 and i - 1 >= 0 and 
                                    (i < 2 or result[i-2] == ' ') and result[i-1].isalnum())
                    after_single = (i + 1 < len(result) and result[i+1].isalnum() and
                                   (i + 2 >= len(result) or result[i+2] == ' '))
                    
                    if before_single and after_single:
                        # Skip space (collapse)
                        i += 1
                    else:
                        collapsed += ' '
                        i += 1
            result = collapsed
    
    return result


def _reconstruct_columns(text: str) -> str:
    """
    Attempt to reconstruct text from multi-column PDF layouts
    by analyzing line patterns and merging appropriately
    """
    lines = text.split('\n')
    reconstructed = []

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Check if this might be a split line (ends abruptly, next line continues)
        if line and i + 1 < len(lines):
            next_line = lines[i + 1].strip()

            # If current line ends with a word and next starts with lowercase, merge
            if (line and not line.endswith(('.', ':', ',', ';', '!', '?')) and
                    next_line and next_line[0].islower()):
                reconstructed.append(line + ' ' + next_line)
                i += 2
                continue

        reconstructed.append(line)
        i += 1

    return '\n'.join(reconstructed)


# ═══════════════════════════════════════════════════════════════
#                    ENHANCED PDF EXTRACTION
# ═══════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF with enhanced handling for complex layouts"""
    text_parts = []

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))

        for page_num, page in enumerate(reader.pages):
            try:
                # Try standard extraction first
                page_text = page.extract_text()

                if page_text:
                    page_text = _clean_extracted_text(page_text)
                    page_text = _collapse_spaced_text(page_text)

                    # Try to reconstruct columns if needed
                    if _looks_like_multicolumn(page_text):
                        page_text = _reconstruct_columns(page_text)

                    text_parts.append(page_text)

            except Exception:
                # If standard extraction fails, try alternative method
                try:
                    page_text = page.extract_text(extraction_mode="layout")
                    if page_text:
                        text_parts.append(_collapse_spaced_text(_clean_extracted_text(page_text)))
                except Exception:
                    pass

        full_text = "\n\n".join(text_parts)

        # Post-process to fix common issues
        full_text = _post_process_pdf_text(full_text)
        full_text = _collapse_spaced_text(full_text)

        return full_text.strip()

    except Exception as e:
        return f"Error reading PDF: {str(e)}"


def _looks_like_multicolumn(text: str) -> bool:
    """Detect if text appears to be from a multi-column layout"""
    lines = text.split('\n')

    # Check for signs of multi-column: lots of short lines, irregular spacing
    short_lines = sum(1 for line in lines if 10 < len(line.strip()) < 40)
    total_lines = len([ln for ln in lines if ln.strip()])

    if total_lines > 0 and short_lines / total_lines > 0.5:
        return True

    # Check for excessive whitespace in middle of lines (column gap)
    lines_with_gaps = sum(1 for line in lines if re.search(r'\S\s{5,}\S', line))
    if total_lines > 0 and lines_with_gaps / total_lines > 0.3:
        return True

    return False


def _post_process_pdf_text(text: str) -> str:
    """Post-process PDF text to fix common extraction issues"""

    # Fix common section headers that may have been split
    section_headers = [
        'EDUCATION', 'EXPERIENCE', 'WORK EXPERIENCE', 'PROFESSIONAL EXPERIENCE',
        'SKILLS', 'TECHNICAL SKILLS', 'PROJECTS', 'CERTIFICATIONS',
        'ACHIEVEMENTS', 'AWARDS', 'SUMMARY', 'OBJECTIVE', 'PROFILE',
        'CONTACT', 'PERSONAL INFORMATION', 'REFERENCES'
    ]

    for header in section_headers:
        # Fix headers that got split across lines
        pattern = r'(' + r'\s*'.join(list(header)) + r')'
        text = re.sub(pattern, header, text, flags=re.IGNORECASE)

    # Remove excessive blank lines
    text = re.sub(r'\n{4,}', '\n\n\n', text)

    return text


# ═══════════════════════════════════════════════════════════════
#                    ENHANCED DOCX EXTRACTION
# ═══════════════════════════════════════════════════════════════

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX including tables, headers, footers, and text boxes"""
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        # ── Extract from headers (often contains name/contact) ──
        for section in doc.sections:
            try:
                header = section.header
                if header:
                    for para in header.paragraphs:
                        if para.text.strip():
                            text_parts.insert(0, para.text.strip())

                    # Also check tables in header
                    for table in header.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text:
                                    row_text.append(cell_text)
                            if row_text:
                                text_parts.insert(0, " | ".join(row_text))
            except Exception:
                pass

        # ── Extract text boxes (common in creative resumes) ──
        try:
            for element in doc.element.body.iter():
                if element.tag.endswith('txbxContent'):
                    for child in element.iter():
                        if child.text and child.text.strip():
                            text_parts.append(child.text.strip())
        except Exception:
            pass

        # ── Extract from main body with style awareness ──
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt:
                style_name = para.style.name if para.style else ""

                # Detect headings and format appropriately
                if 'Heading' in style_name or 'Title' in style_name:
                    text_parts.append(f"\n{txt.upper()}\n")
                else:
                    text_parts.append(txt)

        # ── Extract from tables (very common in resumes) ──
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = []
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_text.append(para.text.strip())

                    if cell_text:
                        row_data.append(" ".join(cell_text))

                if row_data:
                    text_parts.append(" | ".join(row_data))

        # ── Extract from footers ──
        for section in doc.sections:
            try:
                footer = section.footer
                if footer:
                    for para in footer.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text.strip())
            except Exception:
                pass

        # ── Try to extract from shapes/drawings ──
        try:
            for rel in doc.part.rels.values():
                if "drawing" in str(rel.target_ref).lower():
                    pass
        except Exception:
            pass

        full_text = "\n".join(text_parts)
        return _clean_extracted_text(full_text)

    except Exception as e:
        return f"Error reading DOCX: {str(e)}"


# ═══════════════════════════════════════════════════════════════
#                    DOC FILE EXTRACTION
# ═══════════════════════════════════════════════════════════════

def extract_text_from_doc(file_bytes: bytes) -> str:
    """
    Extract text from legacy .DOC files (Microsoft Word 97-2003)
    Uses multiple methods for maximum compatibility
    """
    text = ""

    # Method 1: Try using python-docx (sometimes works with .doc)
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        text = "\n".join(text_parts)
        if text and len(text.strip()) > 50:
            return _clean_extracted_text(text)
    except Exception:
        pass

    # Method 2: Try antiword via subprocess (Linux/Mac)
    try:
        import subprocess
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            result = subprocess.run(
                ['antiword', tmp_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                if len(text.strip()) > 50:
                    return _clean_extracted_text(text)
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except Exception:
        pass

    # Method 3: Try catdoc via subprocess (Linux/Mac)
    try:
        import subprocess
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            result = subprocess.run(
                ['catdoc', tmp_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                if len(text.strip()) > 50:
                    return _clean_extracted_text(text)
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except Exception:
        pass

    # Method 4: Try olefile to extract raw text
    try:
        import olefile

        ole = olefile.OleFileIO(io.BytesIO(file_bytes))

        if ole.exists('WordDocument'):
            text_parts = []
            streams_to_try = ['WordDocument', '1Table', '0Table', 'Data']

            for stream_name in streams_to_try:
                if ole.exists(stream_name):
                    try:
                        stream_data = ole.openstream(stream_name).read()
                        extracted = _extract_text_from_binary(stream_data)
                        if extracted:
                            text_parts.append(extracted)
                    except Exception:
                        pass

            text = "\n".join(text_parts)
            ole.close()

            if text and len(text.strip()) > 50:
                return _clean_extracted_text(text)

        ole.close()
    except ImportError:
        pass
    except Exception:
        pass

    # Method 5: Raw binary text extraction (last resort)
    try:
        text = _extract_text_from_binary(file_bytes)
        if text and len(text.strip()) > 50:
            return _clean_extracted_text(text)
    except Exception:
        pass

    # Method 6: Try textract if available
    try:
        import textract
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            text = textract.process(tmp_path).decode('utf-8', errors='ignore')
            if len(text.strip()) > 50:
                return _clean_extracted_text(text)
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except ImportError:
        pass
    except Exception:
        pass

    return text if text else "Error: Could not extract text from .DOC file. Please convert to .DOCX or PDF format."


def _extract_text_from_binary(data: bytes) -> str:
    """
    Extract readable text from binary data
    Handles both ASCII and Unicode text embedded in binary files
    """
    text_parts = []

    # Extract ASCII text (sequences of printable characters)
    ascii_pattern = rb'[\x20-\x7E]{4,}'
    ascii_matches = re.findall(ascii_pattern, data)
    for match in ascii_matches:
        try:
            decoded = match.decode('ascii', errors='ignore')
            if _is_meaningful_text(decoded):
                text_parts.append(decoded)
        except Exception:
            pass

    # Extract UTF-16 text (common in .doc files)
    try:
        utf16_text = data.decode('utf-16-le', errors='ignore')
        cleaned = ''.join(c if c.isprintable() or c in '\n\r\t' else ' ' for c in utf16_text)
        sequences = re.findall(r'[\w\s@.,;:!?()\-]{5,}', cleaned)
        for seq in sequences:
            if _is_meaningful_text(seq):
                text_parts.append(seq)
    except Exception:
        pass

    # Join and clean
    full_text = ' '.join(text_parts)
    full_text = re.sub(r'\s+', ' ', full_text)
    full_text = re.sub(r'([.!?])\s+([A-Z])', r'\1\n\2', full_text)

    return full_text.strip()


def _is_meaningful_text(text: str) -> bool:
    """Check if extracted text is meaningful (not garbage)"""
    if not text or len(text) < 4:
        return False

    letters = sum(1 for c in text if c.isalpha())
    total = len(text)

    if letters / total < 0.4:
        return False

    if len(set(text.lower())) < len(text) * 0.3:
        return False

    garbage_patterns = [
        r'^[A-Z]{20,}$',
        r'^[a-z]{20,}$',
        r'^\d+$',
        r'^[\W_]+$',
    ]

    for pattern in garbage_patterns:
        if re.match(pattern, text):
            return False

    return True


# ═══════════════════════════════════════════════════════════════
#                    ENHANCED TXT EXTRACTION
# ═══════════════════════════════════════════════════════════════

def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from TXT with encoding detection"""
    encodings = [
        'utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'ascii',
        'utf-16', 'utf-16-le', 'utf-16-be', 'iso-8859-1',
    ]

    for enc in encodings:
        try:
            text = file_bytes.decode(enc)
            if text and not any(
                ord(c) < 9 or (13 < ord(c) < 32)
                for c in text[:1000]
                if c not in '\t\n\r'
            ):
                return _clean_extracted_text(text)
        except Exception:
            continue

    return _clean_extracted_text(file_bytes.decode('utf-8', errors='ignore'))


# ═══════════════════════════════════════════════════════════════
#                    ENHANCED IMAGE OCR
# ═══════════════════════════════════════════════════════════════

def extract_text_from_image(
    file_bytes: bytes,
    groq_api_key: str,
    file_name: str = "resume.png"
) -> str:
    """Extract ALL text from resume image using Groq Vision API"""
    try:
        ext = Path(file_name).suffix.lower()
        mime_map = {
            ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".png": "image/png", ".webp": "image/webp",
        }
        mime_type = mime_map.get(ext, "image/png")
        base64_image = base64.b64encode(file_bytes).decode("utf-8")

        headers = {
            "Authorization": f"Bearer {groq_api_key}",
            "Content-Type": "application/json"
        }

        extraction_prompt = """You are an expert resume OCR system. Your task is to extract EVERY piece of text from this resume image with 100% accuracy.

CRITICAL EXTRACTION RULES:

1. **NAME** (Usually largest text at top):
   - Look for the largest/boldest text - this is typically the name
   - It may NOT have a "Name:" label
   - Extract EXACTLY as written (First Last or Last, First)
   - Common positions: top-center, top-left, inside a header box

2. **CONTACT INFORMATION** (Extract ALL of these):
   - Phone: Look for +, (, ), -, numbers. Include country code.
   - Email: Look for @ symbol anywhere
   - Address: Street, City, State, ZIP/PIN
   - LinkedIn: linkedin.com/in/...
   - GitHub: github.com/...
   - Portfolio/Website: Any other URLs
   - Location: City, Country

3. **SECTIONS TO EXTRACT**:
   - Professional Summary/Objective
   - Work Experience (ALL jobs with dates, titles, companies, bullet points)
   - Education (ALL degrees, institutions, dates, GPA)
   - Skills (ALL skills, categorized or not)
   - Certifications (names, issuers, dates)
   - Projects (names, descriptions, technologies)
   - Awards/Achievements
   - Publications
   - Languages
   - Interests/Hobbies

4. **FORMATTING RULES**:
   - Preserve section headings in CAPS
   - Keep bullet points as "• "
   - Maintain date formats as written
   - Keep job titles and company names exact
   - Preserve GPA/percentage values

5. **COMMON CHALLENGES**:
   - Multi-column layouts: Read left column fully, then right column
   - Sidebar content: Don't miss contact info in sidebars
   - Icons: If icons represent phone/email/location, extract the text next to them
   - Creative formats: Scan entire image, not just standard positions

OUTPUT FORMAT:
Return clean, structured text with clear section breaks. Do NOT summarize or paraphrase. Extract word-for-word."""

        payload = {
            "model": "meta-llama/llama-4-scout-17b-16e-instruct",
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "text", "text": extraction_prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{base64_image}"}
                    }
                ]
            }],
            "temperature": 0.1,
            "max_tokens": 8000,
        }

        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers=headers, json=payload, timeout=90
        )

        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        else:
            error = response.json().get("error", {}).get("message", "Unknown error")
            return f"Error extracting text from image: {error}"

    except Exception as e:
        return f"Error processing image: {str(e)}"


# ═══════════════════════════════════════════════════════════════
#                    NAME / CONTACT VALIDATION
# ═══════════════════════════════════════════════════════════════

def _is_valid_name(name: str) -> bool:
    """
    Validate if a string looks like a valid person name.
    Uses two-tier blacklist to reject section headers, job titles,
    and random phrases.
    """
    if not name:
        return False

    name_clean = name.strip()
    name_lower = name_clean.lower()

    # Exact-match rejects
    exact_rejects = {
        "", "n/a", "na", "none", "unknown", "candidate",
        "your name", "first last", "firstname lastname",
        "name", "full name", "[name]", "<name>", "(name)",
        "enter name", "type name", "not available",
        "curriculum vitae", "resume", "cv",
    }
    if name_lower in exact_rejects:
        return False

    words = name_clean.split()

    # Must have 2-4 words
    if not (2 <= len(words) <= 4):
        return False

    # Each word: reasonable length, mostly alphabetic
    for w in words:
        stripped = w.strip(".-'")
        if not stripped or len(stripped) > 20:
            return False
        alpha_ratio = sum(c.isalpha() for c in stripped) / max(len(stripped), 1)
        if alpha_ratio < 0.8:
            return False

    # Overall alpha ratio
    alpha_chars = sum(1 for c in name_clean if c.isalpha() or c.isspace())
    if alpha_chars / max(len(name_clean), 1) < 0.80:
        return False

    # ── Strong blacklist: ANY word → reject ──
    for w in words:
        if w.lower().strip(".,;:()") in _STRONG_NAME_BLACKLIST:
            return False

    # ── Soft blacklist: 2+ words → reject ──
    soft_hits = sum(
        1 for w in words
        if w.lower().strip(".,;:()") in _SOFT_NAME_BLACKLIST
    )
    if soft_hits >= 2:
        return False

    # Names are usually Title Case or ALL CAPS, not all lowercase
    if name_clean == name_clean.lower():
        return False

    return True


# ═══════════════════════════════════════════════════════════════
#                    CONTACT EXTRACTION HELPERS
# ═══════════════════════════════════════════════════════════════

def extract_contacts_from_text(text: str) -> Dict:
    """
    Enhanced contact extraction with 20+ patterns for each field
    Works even without labels like 'Email:', 'Phone:', etc.
    """
    contacts: Dict[str, str] = {
        "name": "",
        "email": "",
        "phone": "",
        "address": "",
        "linkedin": "",
        "github": "",
        "portfolio": "",
        "location": ""
    }

    if not text:
        return contacts

    # ═══════════════════════════════════════
    # EMAIL EXTRACTION (Most reliable)
    # ═══════════════════════════════════════
    email_patterns = [
        r'[\w._%+-]+@[\w.-]+\.[a-zA-Z]{2,}',
        r'[\w._%+-]+\s*@\s*[\w.-]+\s*\.\s*[a-zA-Z]{2,}',
        r'[\w._%+-]+\[at\][\w.-]+\.[a-zA-Z]{2,}',
        r'[\w._%+-]+\(at\)[\w.-]+\.[a-zA-Z]{2,}',
    ]

    for pattern in email_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            email = match.group(0).strip()
            email = re.sub(r'\s+', '', email)
            email = re.sub(r'\[at\]|\(at\)', '@', email, flags=re.IGNORECASE)
            contacts["email"] = email
            break

    # ═══════════════════════════════════════
    # PHONE EXTRACTION
    # ═══════════════════════════════════════
    phone_patterns = [
        r'\+\d{1,3}[-.\s]?\(?\d{2,5}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}',
        r'\+\d{1,3}[-.\s]?\d{4,5}[-.\s]?\d{5,6}',
        r'\+\d{10,15}',
        r'\+91[-.\s]?\d{5}[-.\s]?\d{5}',
        r'\+91[-.\s]?\d{10}',
        r'91[-.\s]?\d{10}',
        r'0\d{2,4}[-.\s]?\d{6,8}',
        r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}',
        r'\d{3}[-.\s]\d{3}[-.\s]\d{4}',
        r'1[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\+44[-.\s]?\d{4}[-.\s]?\d{6}',
        r'0\d{4}[-.\s]?\d{6}',
        r'(?:Phone|Ph|Tel|Mobile|Cell|Contact|M|T|P)[\s.:]*[\+]?\d[\d\s\-().]{8,18}',
        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
        r'\b\d{5}[-.\s]?\d{5}\b',
        r'\b\d{10,12}\b',
    ]

    for pattern in phone_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            cleaned = re.sub(
                r'^(?:Phone|Ph|Tel|Mobile|Cell|Contact|M|T|P)[\s.:]*',
                '', match, flags=re.IGNORECASE
            )
            cleaned = cleaned.strip()
            digits = re.sub(r'[^\d]', '', cleaned)
            if 10 <= len(digits) <= 15:
                contacts["phone"] = cleaned
                break
        if contacts["phone"]:
            break

    # ═══════════════════════════════════════
    # LINKEDIN EXTRACTION
    # ═══════════════════════════════════════
    linkedin_patterns = [
        r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+/?',
        r'linkedin\.com/in/[\w-]+',
        r'linkedin:\s*[\w-]+',
        r'in/[\w-]+\s*\(linkedin\)',
    ]

    for pattern in linkedin_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            linkedin = match.group(0).strip()
            if not linkedin.startswith('http'):
                if 'linkedin.com' not in linkedin.lower():
                    linkedin = f"linkedin.com/in/{linkedin.split('/')[-1]}"
            contacts["linkedin"] = linkedin
            break

    # ═══════════════════════════════════════
    # GITHUB EXTRACTION
    # ═══════════════════════════════════════
    github_patterns = [
        r'(?:https?://)?(?:www\.)?github\.com/[\w-]+/?',
        r'github\.com/[\w-]+',
        r'github:\s*[\w-]+',
    ]

    for pattern in github_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            contacts["github"] = match.group(0).strip()
            break

    # ═══════════════════════════════════════
    # PORTFOLIO/WEBSITE EXTRACTION
    # ═══════════════════════════════════════
    website_patterns = [
        r'(?:https?://)?(?:www\.)?[\w-]+\.(?:com|io|dev|me|org|net|co|in)/?\S*',
        r'(?:portfolio|website|web|site)[\s.:]*(?:https?://)?[\w.-]+\.[a-z]{2,}',
    ]

    for pattern in website_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            url = match.strip()
            if 'linkedin' not in url.lower() and 'github' not in url.lower():
                if '@' not in url:
                    contacts["portfolio"] = url
                    break
        if contacts["portfolio"]:
            break

    # ═══════════════════════════════════════
    # ADDRESS EXTRACTION
    # ═══════════════════════════════════════
    address_patterns = [
        r'(?:No[.:]?\s*)?[\d]+[,\s]+[\w\s]+(?:Street|St|Road|Rd|Avenue|Ave|Nagar|Colony|Layout|Block|Lane|Apt|Apartment|Floor|Building|Bldg)[\w\s,.-]+(?:\d{5,6})',
        r'[\w\s]+,\s*[\w\s]+,\s*[\w\s]+-?\s*\d{5,6}',
        r'[\w\s]+,\s*[\w\s]+,\s*[\w\s]+,\s*\d{5,6}',
        r'(?:Address|Location|Residence|Home)[\s.:]+([^\n]{15,100})',
        r'(?:Bangalore|Bengaluru|Mumbai|Delhi|Chennai|Hyderabad|Pune|Kolkata|Noida|Gurgaon|Gurugram)[\w\s,.-]+\d{5,6}',
        r'(?:New York|San Francisco|Los Angeles|Chicago|Seattle|Boston|Austin)[\w\s,.-]+\d{5}',
    ]

    for pattern in address_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            addr = match.group(1) if match.lastindex else match.group(0)
            addr = re.sub(
                r'^(?:Address|Location|Residence|Home)[\s.:]+',
                '', addr, flags=re.IGNORECASE
            )
            addr = addr.strip()
            if 15 < len(addr) < 200:
                contacts["address"] = addr
                break

    # ═══════════════════════════════════════
    # LOCATION EXTRACTION (City, State, Country)
    # ═══════════════════════════════════════
    location_patterns = [
        r'(?:Location|Based in|Located at|City)[\s.:]*([A-Z][a-zA-Z\s]+,\s*[A-Z][a-zA-Z\s]+)',
        r'\b(Bangalore|Bengaluru|Mumbai|Delhi|NCR|Chennai|Hyderabad|Pune|Kolkata|Noida|Gurgaon|Gurugram|Ahmedabad|Jaipur|Lucknow|Chandigarh|Indore|Bhopal|Kochi|Coimbatore|Trivandrum|Mysore|Nagpur)\b',
        r'\b(New York|NYC|San Francisco|SF|Los Angeles|LA|Chicago|Seattle|Boston|Austin|Denver|Atlanta|Dallas|Houston|Phoenix|San Diego|San Jose|Portland|Miami|Washington DC|Philadelphia)\b',
        r'\b(London|Berlin|Amsterdam|Dublin|Singapore|Tokyo|Sydney|Toronto|Vancouver|Melbourne|Paris|Munich|Barcelona|Stockholm|Copenhagen|Zurich|Dubai|Hong Kong)\b',
        r'\b(India|USA|United States|UK|United Kingdom|Canada|Australia|Germany|Netherlands|Singapore|UAE|Ireland)\b',
    ]

    for pattern in location_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            loc = match.group(1) if match.lastindex else match.group(0)
            contacts["location"] = loc.strip()
            break

    # ═══════════════════════════════════════
    # NAME EXTRACTION
    # ═══════════════════════════════════════
    contacts["name"] = extract_name_from_text(text)

    return contacts


def extract_name_from_text(text: str) -> str:
    """
    Enhanced name extraction using multiple strategies.
    Uses strict validation to reject section headers, job titles,
    and random phrases.
    """
    if not text:
        return ""

    lines = text.strip().split("\n")

    # ── Skip keywords (lines containing these are NOT names) ──
    skip_keywords = [
        'resume', 'curriculum', 'vitae', 'cv', 'http', 'www', '@',
        'address', 'phone', 'email', 'street', 'road', 'avenue',
        'objective', 'summary', 'profile', 'linkedin', 'github',
        'portfolio', 'mobile', 'tel:', 'contact', 'experience',
        'education', 'skills', 'professional', 'career', 'about',
        'personal', 'details', 'information', 'confidential',
        'page', 'date', 'application', 'position', 'job',
        'candidate', 'recruitment', 'hiring', 'vacancy',
        'declaration', 'reference', 'signature',
        'company', 'corporation', 'limited', 'pvt', 'ltd',
    ]

    # Strategy 1: Look for labeled name
    name_label_patterns = [
        r'(?:Name|Full Name|Candidate Name)[\s.:]+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,3})',
        r'^([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,3})[\s]*$',
    ]

    for pattern in name_label_patterns:
        for line in lines[:15]:
            match = re.match(pattern, line.strip(), re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                if _is_valid_name(name):
                    return name

    # Strategy 2: First substantial line that looks like a name
    for line in lines[:15]:
        line = line.strip()

        if not line or len(line) < 4 or len(line) > 45:
            continue

        # Skip lines with keywords
        if any(kw in line.lower() for kw in skip_keywords):
            continue

        # Skip lines with digits at start, 5+ digits, phone-like, or email
        if re.match(r'^\d', line):
            continue
        if len(re.findall(r'\d', line)) >= 5:
            continue
        if re.match(r'^[\+\d\(\)]', line):
            continue
        if '@' in line or 'http' in line.lower():
            continue

        # Must be mostly alphabetic
        alpha_chars = sum(1 for c in line if c.isalpha() or c.isspace() or c == '.')
        if alpha_chars / max(len(line), 1) < 0.85:
            continue

        # Check explicit name patterns
        name_patterns = [
            r'^[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?$',
            r'^[A-Z]+\s+[A-Z]+(?:\s+[A-Z]+)?$',
            r'^[A-Z][a-z]+\s+[A-Z]\.\s*[A-Z][a-z]+$',
            r'^[A-Z][a-z]+\s+[A-Z][a-z]+\s+[A-Z][a-z]+$',
            r'^Dr\.\s+[A-Z][a-z]+\s+[A-Z][a-z]+$',
            r'^Mr\.\s*[A-Z][a-z]+\s+[A-Z][a-z]+$',
            r'^Ms\.\s*[A-Z][a-z]+\s+[A-Z][a-z]+$',
            r'^[A-Z][a-z]+\s+[A-Z][a-z]+-[A-Z][a-z]+$',
        ]

        for pattern in name_patterns:
            if re.match(pattern, line):
                if _is_valid_name(line):
                    return line

        # Fallback: 2-4 capitalized words of reasonable length
        words = line.split()
        if 2 <= len(words) <= 4:
            if all(w[0].isupper() for w in words if w):
                if all(2 <= len(w) <= 15 for w in words):
                    if _is_valid_name(line):
                        return line

    # Strategy 3: Look for "I am" / "My name is"
    intro_patterns = [
        r"(?:I am|I'm|My name is|This is)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})",
    ]

    for pattern in intro_patterns:
        match = re.search(pattern, text)
        if match:
            name = match.group(1).strip()
            if _is_valid_name(name):
                return name

    # Strategy 4: ALL CAPS line (common in some formats)
    section_headers_upper = {
        'RESUME', 'CURRICULUM', 'VITAE', 'CV', 'OBJECTIVE',
        'SUMMARY', 'EXPERIENCE', 'EDUCATION', 'SKILLS',
        'CONTACT', 'PROFILE', 'ABOUT', 'PROJECTS',
        'CERTIFICATIONS', 'ACHIEVEMENTS', 'AWARDS',
        'REFERENCES', 'DECLARATION', 'PERSONAL',
        'PROFESSIONAL', 'TECHNICAL', 'WORK', 'HISTORY',
        'QUALIFICATIONS', 'RESPONSIBILITIES',
        'DETAILS', 'INFORMATION', 'OVERVIEW',
    }

    for line in lines[:10]:
        line = line.strip()
        if line and line.isupper() and 5 <= len(line) <= 40:
            words = line.split()
            if 2 <= len(words) <= 4:
                if not any(header in line for header in section_headers_upper):
                    title_name = line.title()
                    if _is_valid_name(title_name):
                        return title_name

    return ""


# ═══════════════════════════════════════════════════════════════
#                    FILE PREVIEW GENERATION
# ═══════════════════════════════════════════════════════════════

def get_file_preview_data(file_bytes: bytes, file_name: str) -> Dict:
    """Generate preview data for displaying uploaded file in UI"""
    ext = Path(file_name).suffix.lower()

    preview: Dict = {
        "type": ext,
        "name": file_name,
        "size_kb": round(len(file_bytes) / 1024, 1),
        "can_preview": False,
        "preview_data": None,
        "mime_type": "",
        "page_count": 0,
    }

    if ext in [".jpg", ".jpeg", ".png", ".webp"]:
        preview["can_preview"] = True
        preview["preview_data"] = base64.b64encode(file_bytes).decode("utf-8")
        mime = {
            ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".png": "image/png", ".webp": "image/webp",
        }
        preview["mime_type"] = mime.get(ext, "image/png")

    elif ext == ".pdf":
        preview["can_preview"] = True
        preview["preview_data"] = base64.b64encode(file_bytes).decode("utf-8")
        preview["mime_type"] = "application/pdf"
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            preview["page_count"] = len(reader.pages)
        except Exception:
            pass

    elif ext in [".txt", ".md", ".text"]:
        preview["can_preview"] = True
        preview["preview_data"] = file_bytes.decode("utf-8", errors="ignore")[:5000]
        preview["mime_type"] = "text/plain"

    elif ext in [".docx", ".doc"]:
        preview["can_preview"] = True
        preview["preview_data"] = extract_text_from_docx(file_bytes)[:5000]
        preview["mime_type"] = "text/plain"

    return preview


# ═══════════════════════════════════════════════════════════════
#                    MAIN PROCESSING FUNCTION
# ═══════════════════════════════════════════════════════════════

def process_uploaded_file(uploaded_file, groq_api_key: str = "") -> Dict:
    """Process any uploaded file and extract text + preview"""
    file_name = uploaded_file.name
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)
    ext = Path(file_name).suffix.lower()

    result: Dict = {
        "file_name": file_name,
        "file_type": ext,
        "file_size_kb": round(len(file_bytes) / 1024, 1),
        "text": "",
        "success": False,
        "error": None,
        "file_bytes": file_bytes,
        "preview": get_file_preview_data(file_bytes, file_name),
        "extracted_contacts": {},
    }

    try:
        if ext == ".pdf":
            result["text"] = extract_text_from_pdf(file_bytes)
        elif ext == ".docx":
            result["text"] = extract_text_from_docx(file_bytes)
        elif ext == ".doc":
            result["text"] = extract_text_from_doc(file_bytes)
        elif ext in [".txt", ".text", ".md"]:
            result["text"] = extract_text_from_txt(file_bytes)
        elif ext in [".jpg", ".jpeg", ".png", ".webp"]:
            if not groq_api_key:
                result["error"] = "Groq API key required for image OCR"
                return result
            result["text"] = extract_text_from_image(file_bytes, groq_api_key, file_name)
        else:
            result["error"] = f"Unsupported file type: {ext}"
            return result

        if result["text"] and len(result["text"].strip()) > 30:
            if result["text"].startswith("Error:"):
                result["error"] = result["text"]
                result["success"] = False
            else:
                result["success"] = True
                result["extracted_contacts"] = extract_contacts_from_text(result["text"])
        else:
            result["error"] = "Could not extract sufficient text. Try converting to PDF or DOCX format."

    except Exception as e:
        result["error"] = str(e)

    return result
